#!/usr/bin/env python3
"""Generate JAMA-style Word tables for manuscript and supplement.

Outputs:
    outputs/docx/manuscript_tables.docx  — Tables 1-3
    outputs/docx/supplement_tables.docx  — eTables 1-8

Usage: uv run scripts/generate_docx_tables.py
"""

from __future__ import annotations

import re
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import polars as pl
from docx import Document
from docx.document import Document as DocumentType
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

# ── Paths ────────────────────────────────────────────────────────────────
ROOT = Path(__file__).resolve().parent.parent
TABLES_DIR = ROOT / "outputs" / "tables"
OUT_DIR = ROOT / "outputs" / "docx"

# ── JAMA constants ───────────────────────────────────────────────────────
FONT = "Arial"
BODY_PT = 10
TITLE_PT = 10
FOOTNOTE_PT = 9

# Borders
THIN_BORDER = {"sz": "4", "val": "single", "color": "000000"}
NO_BORDER = {"sz": "0", "val": "none", "color": "auto"}

# US Letter (twips: 1 inch = 1440 twips)
PAGE_W = Inches(8.5)
PAGE_H = Inches(11)
MARGIN = Inches(1)


# ── XML helpers ──────────────────────────────────────────────────────────
def set_cell_borders(cell: Any, **kwargs: dict[str, str]) -> None:
    """Set borders on a table cell via XML.

    Usage: set_cell_borders(cell, top=THIN_BORDER, bottom=THIN_BORDER, left=NO_BORDER, right=NO_BORDER)
    """
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge, attrs in kwargs.items():
        el = OxmlElement(f"w:{edge}")
        for k, v in attrs.items():
            el.set(qn(f"w:{k}"), str(v))
        borders.append(el)
    # Remove existing borders before appending
    for existing in tc_pr.findall(qn("w:tcBorders")):
        tc_pr.remove(existing)
    tc_pr.append(borders)


def set_cell_margins(
    cell: Any, top: int = 30, bottom: int = 30, left: int = 60, right: int = 60
) -> None:
    """Set cell margins in twips via XML."""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("start", left), ("end", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        margins.append(el)
    tc_pr.append(margins)


def set_table_width_100pct(table: Any) -> None:
    """Force table width to 100% of page content area via XML."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:type"), "pct")
    tbl_w.set(qn("w:w"), "5000")  # 5000 = 100%
    # Remove existing tblW if present
    for existing in tbl_pr.findall(qn("w:tblW")):
        tbl_pr.remove(existing)
    tbl_pr.append(tbl_w)


def set_row_spacing(cell: Any, line: int = 240) -> None:
    """Set exact line spacing on the paragraph inside a cell."""
    for para in cell.paragraphs:
        pf = para.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        # Exact line spacing via XML
        p_pr = para._p.get_or_add_pPr()
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:line"), str(line))
        spacing.set(qn("w:lineRule"), "exact")
        # Remove existing spacing
        for existing in p_pr.findall(qn("w:spacing")):
            p_pr.remove(existing)
        p_pr.append(spacing)


# ── Helpers ──────────────────────────────────────────────────────────────
def is_numeric_cell(val: str) -> bool:
    """Detect if a cell value should be right-aligned (numeric content)."""
    if not val or val in ("—", "", "N/A"):
        return False
    stripped = re.sub(r"[,()%]", "", val).strip()
    return bool(re.match(r"^-?\d", stripped))


def load_csv(filename: str) -> tuple[list[str], list[list[str]]] | None:
    """Read a CSV file via Polars and return (headers, data_rows).

    Returns None if the file doesn't exist.
    """
    filepath = TABLES_DIR / filename
    if not filepath.exists():
        print(f"WARNING: {filename} not found, skipping.", file=sys.stderr)
        return None
    df = pl.read_csv(filepath, infer_schema_length=0)  # all strings
    headers = df.columns
    data = [list(row) for row in df.iter_rows()]
    return headers, data


# ── Table construction ───────────────────────────────────────────────────
def build_table(doc: DocumentType, headers: list[str], data: list[list[str]]) -> None:
    """Build a JAMA three-line table in the document."""
    n_rows = len(data) + 1  # +1 for header
    n_cols = len(headers)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width_100pct(table)

    # Remove default table borders
    tbl_pr = table._tbl.tblPr
    tbl_borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:color"), "auto")
        tbl_borders.append(el)
    # Remove existing borders
    for existing in tbl_pr.findall(qn("w:tblBorders")):
        tbl_pr.remove(existing)
    tbl_pr.append(tbl_borders)

    # Header row
    header_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ""  # clear default
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(h)
        run.bold = True
        run.font.name = FONT
        run.font.size = Pt(BODY_PT)
        set_cell_borders(cell, top=THIN_BORDER, bottom=THIN_BORDER, left=NO_BORDER, right=NO_BORDER)
        set_cell_margins(cell, top=40, bottom=40, left=60, right=60)
        set_row_spacing(cell)

    # Data rows
    for row_idx, row_data in enumerate(data):
        is_last = row_idx == len(data) - 1
        table_row = table.rows[row_idx + 1]
        for col_idx, val in enumerate(row_data):
            cell = table_row.cells[col_idx]
            cell.text = ""
            para = cell.paragraphs[0]
            numeric = is_numeric_cell(val)
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT if numeric else WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run(val if val is not None else "")
            run.font.name = FONT
            run.font.size = Pt(BODY_PT)
            bottom = THIN_BORDER if is_last else NO_BORDER
            set_cell_borders(cell, top=NO_BORDER, bottom=bottom, left=NO_BORDER, right=NO_BORDER)
            set_cell_margins(cell, top=30, bottom=30, left=60, right=60)
            set_row_spacing(cell)


def add_table_title(doc: DocumentType, label: str, title: str) -> None:
    """Add 'Table N. Title' paragraph."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    run_label = para.add_run(f"{label}. ")
    run_label.bold = True
    run_label.font.name = FONT
    run_label.font.size = Pt(TITLE_PT)
    run_title = para.add_run(title)
    run_title.font.name = FONT
    run_title.font.size = Pt(TITLE_PT)


def add_footnotes(doc: DocumentType, *lines: str) -> None:
    """Add footnote paragraphs (9pt)."""
    for i, line in enumerate(lines):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(4) if i == 0 else Pt(1)
        para.paragraph_format.space_after = Pt(1)
        run = para.add_run(line)
        run.font.name = FONT
        run.font.size = Pt(FOOTNOTE_PT)


def set_section_landscape(section: Any) -> None:
    """Configure a section for landscape orientation."""
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = PAGE_H  # swap
    section.page_height = PAGE_W
    section.top_margin = MARGIN
    section.bottom_margin = MARGIN
    section.left_margin = MARGIN
    section.right_margin = MARGIN


def set_section_portrait(section: Any) -> None:
    """Configure a section for portrait orientation."""
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = PAGE_W
    section.page_height = PAGE_H
    section.top_margin = MARGIN
    section.bottom_margin = MARGIN
    section.left_margin = MARGIN
    section.right_margin = MARGIN


# ── Table definitions ────────────────────────────────────────────────────
@dataclass
class TableDef:
    csv_file: str
    label: str
    title: str
    footnotes: list[str] = field(default_factory=list)
    landscape: bool = False


MANUSCRIPT_TABLES = [
    TableDef(
        csv_file="table1.csv",
        label="Table 1",
        title="Patient Characteristics by Fall Status",
        footnotes=[
            "Abbreviations: SD, standard deviation; IQR, interquartile range; "
            "SMD, standardized mean difference; PMFRS, Predictive Model Fall Risk Score.",
            "Continuous variables reported as mean \u00b1 SD or median [IQR]. "
            "Categorical variables reported as n (%).",
            "P values from Wilcoxon rank-sum test (continuous) or chi-square test (categorical).",
        ],
    ),
    TableDef(
        csv_file="table2.csv",
        label="Table 2",
        title="Discrimination Performance: Epic PMFRS vs Morse Fall Scale",
        landscape=True,
        footnotes=[
            "Abbreviations: AUROC, area under the receiver operating characteristic curve; "
            "AUPRC, area under the precision-recall curve; PPV, positive predictive value; "
            "NPV, negative predictive value; NNE, number needed to evaluate; CI, confidence interval.",
            "DeLong p = paired DeLong test for AUROC comparison between models.",
            "Bootstrap 95% CIs from 2000 stratified resamples (BCa method, seed = 42).",
        ],
    ),
    TableDef(
        csv_file="table3.csv",
        label="Table 3",
        title="Reclassification Analysis (NRI, IDI)",
        footnotes=[
            "Abbreviations: NRI, net reclassification improvement; "
            "IDI, integrated discrimination improvement; CI, confidence interval.",
            "Continuous NRI and categorical NRI reported separately per Pepe et al (Stat Med, 2015).",
            "Event and non-event NRI reported independently. "
            "Positive values favor the Morse Fall Scale over Epic PMFRS.",
            "95% CIs from 2000 stratified bootstrap resamples.",
        ],
    ),
]

SUPPLEMENT_TABLES = [
    TableDef(
        csv_file="etable1_age.csv",
        label="eTable 1",
        title="Discrimination by Age Group",
        footnotes=[
            "Abbreviations: AUROC, area under the receiver operating characteristic curve; "
            "CI, confidence interval; PMFRS, Predictive Model Fall Risk Score.",
            "95% CIs from 2000 stratified bootstrap resamples (BCa method).",
            "Subgroups with fewer than 20 fall events are flagged as unreliable.",
        ],
    ),
    TableDef(
        csv_file="etable2_race.csv",
        label="eTable 2",
        title="Discrimination by Race/Ethnicity",
        footnotes=[
            "Abbreviations: AUROC, area under the receiver operating characteristic curve; "
            "CI, confidence interval; PMFRS, Predictive Model Fall Risk Score.",
            "95% CIs from 2000 stratified bootstrap resamples (BCa method).",
            "Subgroups with fewer than 20 events marked with em-dash (\u2014); AUROC not computed.",
        ],
    ),
    TableDef(
        csv_file="etable3_unit.csv",
        label="eTable 3",
        title="Discrimination by Admitting Department",
        landscape=True,
        footnotes=[
            "Abbreviations: AUROC, area under the receiver operating characteristic curve; "
            "CI, confidence interval; IMCU, intermediate care unit; MICU, medical intensive care unit; "
            "ICU, intensive care unit; PMFRS, Predictive Model Fall Risk Score.",
            "Top 10 departments by encounter volume shown.",
            "95% CIs from 2000 stratified bootstrap resamples (BCa method).",
            "Subgroups with fewer than 20 fall events are flagged as unreliable.",
        ],
    ),
    TableDef(
        csv_file="etable4_sensitivity.csv",
        label="eTable 4",
        title="Sensitivity Analyses Across Score Timings",
        landscape=True,
        footnotes=[
            "Abbreviations: AUROC, area under the receiver operating characteristic curve; "
            "CI, confidence interval; PMFRS, Predictive Model Fall Risk Score.",
            "DeLong p = paired DeLong test comparing Epic PMFRS vs Morse Fall Scale "
            "within each timing strategy.",
            '"Before fall" analysis uses last pre-discharge score for non-fallers as comparator.',
            "Maximum and mean encounter scores include post-fall assessments, "
            "introducing look-ahead bias.",
        ],
    ),
    TableDef(
        csv_file="calibration_summary.csv",
        label="eTable 5",
        title="Calibration Summary: Epic PMFRS and Morse Fall Scale",
        footnotes=[
            "Abbreviations: CITL, calibration-in-the-large; ICI, integrated calibration index; "
            "E:O, expected-to-observed ratio; CI, confidence interval; "
            "PMFRS, Predictive Model Fall Risk Score.",
            "Calibration assessed at admission using logistic recalibration.",
            "ICI = mean absolute difference between LOWESS-smoothed observed and predicted probabilities.",
        ],
    ),
    TableDef(
        csv_file="threshold_summary.csv",
        label="eTable 6",
        title="Threshold Analysis: Optimal Cutpoints by Method",
        landscape=True,
        footnotes=[
            "Abbreviations: PPV, positive predictive value; NPV, negative predictive value; "
            "NNE, number needed to evaluate; NMB, net monetary benefit; "
            "DCA, decision curve analysis; CI, confidence interval; "
            "PMFRS, Predictive Model Fall Risk Score.",
            "Threshold methods: Youden index, closest-to-(0,1), fixed sensitivity (60%, 80%), "
            "value-optimizing (NMB with Monte Carlo), and DCA-derived.",
            "Standard cutoffs: MFS \u226525 (moderate), \u226545 (high risk); "
            "Epic 3-tier (35/70), 2-tier (50).",
        ],
    ),
    TableDef(
        csv_file="figure3_dca_net_benefit.csv",
        label="eTable 7",
        title="Decision Curve Analysis: Net Benefit at Selected Threshold Probabilities",
        footnotes=[
            "Abbreviations: PMFRS, Predictive Model Fall Risk Score; "
            "MFS, Morse Fall Scale; CI, confidence interval.",
            "Net benefit from decision curve analysis (Vickers et al, Diagn Progn Res, 2019).",
            "Threshold probability range 0%\u201310% selected for clinical relevance "
            "given 1.27% fall prevalence.",
        ],
    ),
    TableDef(
        csv_file="etable4_gender.csv",
        label="eTable 8",
        title="Discrimination by Gender",
        footnotes=[
            "Abbreviations: AUROC, area under the receiver operating characteristic curve; "
            "CI, confidence interval; PMFRS, Predictive Model Fall Risk Score.",
            "95% CIs from 2000 stratified bootstrap resamples (BCa method).",
            "Subgroups with fewer than 20 fall events are flagged as unreliable.",
        ],
    ),
]


# ── Document builder ─────────────────────────────────────────────────────
def build_document(table_defs: list[TableDef], out_path: Path):
    """Build a DOCX file from a list of table definitions."""
    doc = Document()

    # Set default font (python-docx stubs type this as BaseStyle; the runtime
    # object is _ParagraphStyle which has .font — safe to ignore the type error)
    style = doc.styles["Normal"]
    style.font.name = FONT  # type: ignore[union-attr]
    style.font.size = Pt(BODY_PT)  # type: ignore[union-attr]

    first_table = True
    for tdef in table_defs:
        result = load_csv(tdef.csv_file)
        if result is None:
            continue

        headers, data = result

        if first_table:
            # Use the default first section
            section = doc.sections[0]
            first_table = False
        else:
            # Add a new section (page break)
            section = doc.add_section(WD_SECTION.NEW_PAGE)

        if tdef.landscape:
            set_section_landscape(section)
        else:
            set_section_portrait(section)

        add_table_title(doc, tdef.label, tdef.title)
        build_table(doc, headers, data)
        add_footnotes(doc, *tdef.footnotes)

    if first_table:
        # No tables were added (all CSVs missing)
        print(f"WARNING: No tables found for {out_path.name}, skipping.", file=sys.stderr)
        return

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    size_kb = out_path.stat().st_size / 1024
    print(f"CREATED: {out_path} ({size_kb:.1f} KB)")


def main() -> None:
    build_document(MANUSCRIPT_TABLES, OUT_DIR / "manuscript_tables.docx")
    build_document(SUPPLEMENT_TABLES, OUT_DIR / "supplement_tables.docx")
    print("\nDone. Both DOCX files generated.")


if __name__ == "__main__":
    main()
